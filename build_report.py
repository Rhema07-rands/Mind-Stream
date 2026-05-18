"""
Script to combine all project report markdown files into a single .docx file.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# -- Default style --
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

def add_heading_tnr(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_centered(text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_list_item(text):
    p = doc.add_paragraph(style='List Paragraph')
    # Clear default and add custom run
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# ============================================================
# READ ALL SOURCE FILES
# ============================================================
files = [
    r'c:\Users\HP\Desktop\MindStream\Project_Report_Chapter1.md',
    r'c:\Users\HP\Desktop\MindStream\Project_Report_Chapter2.md',
    r'c:\Users\HP\Desktop\MindStream\Project_Report_Chapter3.md',
    r'c:\Users\HP\Desktop\MindStream\Project_Report_Chapters4_5.md',
]

all_lines = []
for f in files:
    with open(f, 'r', encoding='utf-8') as fh:
        all_lines.extend(fh.readlines())
        all_lines.append('\n')

# ============================================================
# PARSE AND BUILD DOCUMENT
# ============================================================
i = 0
in_table = False
table_rows = []

def flush_table():
    global table_rows, in_table
    if not table_rows:
        in_table = False
        return
    # Determine columns from first row
    cols = len(table_rows[0])
    tbl = doc.add_table(rows=len(table_rows), cols=cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_data in enumerate(table_rows):
        for ci in range(cols):
            cell_text = row_data[ci] if ci < len(row_data) else ''
            cell = tbl.cell(ri, ci)
            cell.text = cell_text.strip()
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                if ri == 0:
                    for run in paragraph.runs:
                        run.bold = True
    table_rows = []
    in_table = False
    doc.add_paragraph()  # spacing after table

while i < len(all_lines):
    line = all_lines[i].rstrip('\n')

    # Skip markdown horizontal rules
    if line.strip() == '---':
        if in_table:
            flush_table()
        doc.add_page_break()
        i += 1
        continue

    # Blank line
    if line.strip() == '':
        if in_table:
            flush_table()
        i += 1
        continue

    # Table rows (pipe-delimited) — any line with | and 2+ resulting cells
    if '|' in line:
        cells = [c.strip() for c in line.split('|')]
        # Remove empty first/last from leading/trailing pipes
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        # Skip separator rows like "---|---|---"
        if all(re.match(r'^[-:]+$', c) for c in cells):
            i += 1
            continue
        if len(cells) >= 2:
            in_table = True
            table_rows.append(cells)
            i += 1
            continue

    if in_table:
        flush_table()

    # Detect centered title-like lines (all caps, short)
    stripped = line.strip()

    # [PLACEHOLDER: ...] lines
    if stripped.startswith('[PLACEHOLDER'):
        add_para(stripped, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue

    # Figure captions
    if stripped.startswith('Fig ') and ':' in stripped:
        add_para(stripped, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        i += 1
        continue

    # Headings
    if stripped.startswith('CHAPTER '):
        if in_table:
            flush_table()
        add_heading_tnr(stripped, level=1)
        i += 1
        continue

    if stripped.startswith('APPENDIX ') or stripped.startswith('APPENDICES'):
        add_heading_tnr(stripped, level=1)
        i += 1
        continue

    if stripped == 'REFERENCES':
        add_heading_tnr('REFERENCES', level=1)
        i += 1
        continue

    # Section headings like "1.1 Background..."
    m = re.match(r'^(\d+\.\d+(?:\.\d+)?)\s+(.+)$', stripped)
    if m:
        add_heading_tnr(stripped, level=2)
        i += 1
        continue

    # Sub-section bold labels
    if stripped in ['DECLARATION', 'CERTIFICATION', 'APPROVAL', 'DEDICATION',
                    'ACKNOWLEDGMENTS', 'ABSTRACT', 'TABLE OF CONTENTS',
                    'LIST OF FIGURES', 'LIST OF TABLES',
                    'Input Specifications', 'Output Specifications',
                    'System Results', 'Discussion of Findings',
                    'Design Algorithm for Resource Upload Workflow:',
                    'Development Environment:', 'Server/Production Environment:',
                    'End-User Requirements:', 'Development Software:',
                    'Production Software:', 'End-User Software:',
                    'Scoring Guide (For Researcher Reference)',
                    'The Identified Research Gap',
                    'Digital Library Foundations and Early E-Learning Systems',
                    'Learning Management Systems and E-Learning Platforms',
                    'Nigerian Academic Context and ICT Adoption',
                    'Messaging Platforms in Education',
                    'Web Application Technologies',
                    'Database Systems and Cloud Infrastructure',
                    'Security, Authentication, and Access Control',
                    'Usability Evaluation and Human-Computer Interaction',
                    'File Processing and Content Management',
                    'Responsive Design and Mobile Web',
                    'Content Distribution and Access Equity',
                    'Human Usability and System Acceptance:']:
        add_heading_tnr(stripped, level=2)
        i += 1
        continue

    # Phase headings
    if stripped.startswith('Phase '):
        add_heading_tnr(stripped, level=2)
        i += 1
        continue

    # Centered uppercase blocks (title pages etc)
    if stripped.isupper() and len(stripped) < 120:
        add_centered(stripped, bold=True)
        i += 1
        continue

    # Roman numeral list items
    if re.match(r'^(i{1,3}|iv|vi{0,3})\.\s', stripped):
        add_list_item(stripped)
        i += 1
        continue

    # Numbered list items from algorithm
    if re.match(r'^\d+\.\s', stripped) and not re.match(r'^\d+\.\d+', stripped):
        add_list_item(stripped)
        i += 1
        continue

    # Tab-delimited lines (table of contents style)
    if '\t' in stripped:
        add_para(stripped)
        i += 1
        continue

    # START / END keywords
    if stripped in ['START', 'END']:
        add_para(stripped, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1
        continue

    # Note: lines
    if stripped.startswith('Note:'):
        add_para(stripped, italic=True)
        i += 1
        continue

    # Default body text
    add_para(stripped)
    i += 1

# Flush any remaining table
if in_table:
    flush_table()

# ============================================================
# SAVE
# ============================================================
output_path = r'c:\Users\HP\Desktop\MindStream\MindStream_Project_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
